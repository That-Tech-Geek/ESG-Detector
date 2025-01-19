import streamlit as st
import yfinance as yf
import numpy as np
import pandas as pd
from scipy.optimize import minimize
from datetime import datetime
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Function to download and clean data
def download_and_clean_data(tickers, period="max", interval="1d"):
    data_dict = {}
    failed_tickers = []
    success_tickers = []

    for ticker in tickers:
        try:
            data = yf.download(ticker, period=period, interval=interval)['Adj Close']
            if data.isna().sum().sum() == 0:
                data_dict[ticker] = data
                success_tickers.append(ticker)
            else:
                failed_tickers.append(ticker)
        except Exception as e:
            failed_tickers.append(ticker)

    if not data_dict:
        return None, []

    data = pd.DataFrame(data_dict)
    return data, list(data.columns)

# Function to calculate portfolio performance
def portfolio_performance(weights, returns, cov_matrix):
    portfolio_return = np.sum(returns.mean() * weights) * 252
    portfolio_stddev = np.sqrt(np.dot(weights.T, np.dot(cov_matrix, weights))) * np.sqrt(252)
    return portfolio_return, portfolio_stddev

# Sharpe ratio optimization
def negative_sharpe_ratio(weights, returns, cov_matrix, risk_free_rate=0.03):
    p_return, p_stddev = portfolio_performance(weights, returns, cov_matrix)
    sharpe_ratio = (p_return - risk_free_rate) / p_stddev
    return -sharpe_ratio

# Maximize Sharpe Ratio
def maximize_sharpe_ratio(data):
    returns = data.pct_change(fill_method=None).dropna()
    cov_matrix = returns.cov()
    num_assets = len(data.columns)
    initial_guess = num_assets * [1. / num_assets]
    bounds = tuple((0, 1) for _ in range(num_assets))
    constraints = {'type': 'eq', 'fun': lambda x: np.sum(x) - 1}

    result = minimize(negative_sharpe_ratio, initial_guess, args=(returns, cov_matrix),
                      method='SLSQP', bounds=bounds, constraints=constraints)

    if not result.success:
        return None, None, None, None

    optimal_weights = result.x
    optimal_return, optimal_stddev = portfolio_performance(optimal_weights, returns, cov_matrix)
    optimal_sharpe = -result.fun
    return optimal_weights, optimal_sharpe, optimal_return, optimal_stddev

# Generate report
def generate_docx_report(date, tickers, weights, sharpe, return_, stddev):
    doc = Document()
    doc.add_heading(f"Portfolio Report for {date}", level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_heading("Portfolio Summary", level=2)
    doc.add_paragraph(f"Maximized Sharpe Ratio: {sharpe:.4f}")
    doc.add_paragraph(f"Annualized Portfolio Return: {return_:.4%}")
    doc.add_paragraph(f"Annualized Portfolio Volatility: {stddev:.4%}")

    doc.add_heading("Optimal Weights by Ticker", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Ticker"
    header_cells[1].text = "Weight"
    for ticker, weight in zip(tickers, weights):
        row_cells = table.add_row().cells
        row_cells[0].text = ticker
        row_cells[1].text = f"{weight:.4%}"

    return doc

# Streamlit App
st.title("Portfolio Optimization App")
st.sidebar.header("Configuration")

# Ticker Selection
selected_tickers = st.sidebar.multiselect("Select Tickers", options=tickers, default=tickers[:5])
if not selected_tickers:
    st.error("Please select at least one ticker.")

# Download and Clean Data
if st.sidebar.button("Optimize Portfolio"):
    with st.spinner("Downloading and processing data..."):
        data, valid_tickers = download_and_clean_data(selected_tickers)
        if data is None or data.empty:
            st.error("No valid data available for the selected tickers.")
        else:
            optimal_weights, optimal_sharpe, optimal_return, optimal_stddev = maximize_sharpe_ratio(data)
            if optimal_weights is None:
                st.error("Portfolio optimization failed.")
            else:
                # Display results
                st.success("Portfolio optimization complete!")
                st.subheader("Optimization Results")
                st.write(f"**Maximized Sharpe Ratio:** {optimal_sharpe:.4f}")
                st.write(f"**Annualized Return:** {optimal_return:.4%}")
                st.write(f"**Annualized Volatility:** {optimal_stddev:.4%}")

                # Show optimal weights
                weights_df = pd.DataFrame({
                    "Ticker": valid_tickers,
                    "Weight": [f"{w:.4%}" for w in optimal_weights]
                })
                st.write("**Optimal Weights:**")
                st.dataframe(weights_df)

                # Generate and download report
                date = datetime.now().strftime("%Y-%m-%d")
                doc = generate_docx_report(date, valid_tickers, optimal_weights, optimal_sharpe, optimal_return, optimal_stddev)

                # Save the report to memory and provide download link
                st.download_button(
                    label="Download Report",
                    data=doc_to_bytes(doc),
                    file_name=f"Portfolio_Report_{date}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
else:
    st.info("Configure settings and click 'Optimize Portfolio' to begin.")

# Helper function to convert docx to bytes
def doc_to_bytes(doc):
    from io import BytesIO
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()
